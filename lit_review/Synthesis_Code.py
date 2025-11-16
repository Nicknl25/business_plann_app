# ENTRYPOINT: flask_app
# PURPOSE: Launches a web interface for Literature Synthesis generation.
# HOW TO RUN:  python lit_review/Synthesis_Code.py

import os, re, json, pandas as pd
from flask import Flask, render_template_string, request, jsonify
from threading import Lock, Thread, Event
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from datetime import datetime
import xlsxwriter

# ---------- SETUP ----------
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
app = Flask(__name__)

_progress_lock = Lock()
cancel_event = Event()
progress_state = {"active": False, "done": False, "message": "", "log": [], "current": 0, "total": 0}
progress_thread = None

themes_fixed = [
    "Forecast Accuracy Across Methods",
    "Applications in Financial Forecasting",
    "Model Interpretability and Transparency",
    "Firm Size and Contextual Factors",
    "Evaluation Metrics and Methodology",
    "Limitations and Research Gaps",
    "AI versus Other Forecasting Methods"
]

# ---------- FRONTEND ----------
HTML = """
<!DOCTYPE html>
<html>
<head>
<title>Literature Synthesis Tool</title>
<style>
body {font-family:Arial;margin:40px;}
label{display:block;margin-top:12px;font-weight:bold;}
input[type=text]{width:80%;padding:8px;}
button{margin-top:20px;padding:10px 20px;background:#0078D4;color:white;border:none;border-radius:5px;cursor:pointer;}
button:hover{background:#005ea0;}
.progress-wrap{margin-top:24px;max-width:800px;}
.bar{width:100%;background:#eee;border-radius:6px;overflow:hidden;height:16px;}
.bar>span{display:block;height:100%;background:#4caf50;width:0%;transition:width .25s ease;}
.log{margin-top:12px;background:#111;color:#0f0;padding:10px;border-radius:6px;height:180px;overflow:auto;font-family:Consolas,monospace;font-size:12px;}
.hide{display:none;}
.success{color:#0a7d00;font-weight:bold;}
.cancel-btn{margin-left:10px;padding:6px 12px;background:#c62828;color:#fff;border:none;border-radius:4px;cursor:pointer;}
.cancel-btn:hover{background:#a61e1e;}
</style>
</head>
<body>
<h2>dY's Literature Synthesis Tool</h2>
<form id="synthForm" method="POST">
  <label>Lit Review Excel File Path:</label>
  <input type="text" name="lit_file" required>
  <label>Output Excel File Path (.xlsx):</label>
  <input type="text" name="output_file" required>
  <button type="submit">Run Synthesis</button>
</form>

<div id="liveProgress" class="progress-wrap hide">
  <div class="bar"><span id="barFill"></span></div>
  <button id="cancelBtn" class="cancel-btn hide" type="button">Cancel</button>
  <div class="log" id="log"></div>
  <div id="finalMsg" class="success hide"></div>
</div>

<script>
window.addEventListener("DOMContentLoaded",()=>{
 const form=document.getElementById('synthForm');
 const live=document.getElementById('liveProgress');
 const bar=document.getElementById('barFill');
 const logBox=document.getElementById('log');
 const finalMsg=document.getElementById('finalMsg');
 const cancelBtn=document.getElementById('cancelBtn');
 let poller=null;

 function startPoll(){
   poller=setInterval(async()=>{
     const r=await fetch('/progress');const p=await r.json();
     if(p.total>0){bar.style.width=Math.min(100,Math.round(p.current/p.total*100))+'%';}
     if(p.log){logBox.textContent=p.log.join('\\n');logBox.scrollTop=logBox.scrollHeight;}
     if(p.done){clearInterval(poller);poller=null;cancelBtn.classList.add('hide');
       finalMsg.textContent=p.message;finalMsg.classList.remove('hide');}
   },1000);
 }

 form.addEventListener('submit',async e=>{
   e.preventDefault();live.classList.remove('hide');cancelBtn.classList.remove('hide');
   finalMsg.classList.add('hide');logBox.textContent='';bar.style.width='0%';
   const fd=new FormData(form);
   await fetch('/',{method:'POST',body:fd});
   startPoll();
 });

 cancelBtn.onclick=async()=>{await fetch('/cancel',{method:'POST'});cancelBtn.disabled=true;};
});
</script>
</body>
</html>
"""

# ---------- BACKEND ----------
def run_synthesis(lit_file, output_file):
    try:
        global cancel_event
        with _progress_lock:
            progress_state.update({
                "active": True,
                "done": False,
                "message": "Loading data...",
                "log": ["Loading data..."],
                "current": 0,
                "total": 5
            })

        if cancel_event.is_set(): return
        lit_sheets = pd.read_excel(lit_file, sheet_name=None)
        lit_df = pd.concat(lit_sheets.values(), ignore_index=True)

        records=[]
        for _,r in lit_df.iterrows():
            methods=str(r.get("Methods",""))
            full=" | ".join(str(r.get(c,"")) for c in ["Overall","Methods","Conclusion"] if pd.notna(r.get(c,"")))
            if methods.strip() or full.strip():
                records.append({"title":str(r.get("Title","")), "apa":str(r.get("APA Citation","")),
                                "file":str(r.get("Source File","")), "methods":methods, "full":full})
        if not records: raise Exception("⚠️ No records found with notes.")

        all_methods="\n".join(r["methods"] for r in records if r["methods"].strip())
        all_full="\n".join(r["full"] for r in records if r["full"].strip())

        prompt=f"""
        You are assisting with a dissertation literature review.
        Research Question: How does ChatGPT’s ability to forecast EBIT compare to traditional forecasting methods
        (regression and ARIMA) for small-cap, mid-cap, and large-cap firms?

        Methods notes: {all_methods}
        Full notes: {all_full}

        Tasks:
        1. Assign every article to one of these seven themes: {themes_fixed}
        2. Identify at least 5 distinct research gaps.
        3. Write one synthesis paragraph per theme.
        4. Write one Integrative paragraph connecting them.
        5. Provide APA reference list.
        Output valid JSON with keys: themes, gaps, synthesis, references.
        """

        with _progress_lock:
            progress_state["message"] = "Contacting GPT-5..."
            progress_state["log"].append("Contacting GPT-5...")
            progress_state["current"] = 1
        if cancel_event.is_set(): return

        response=client.chat.completions.create(model="gpt-5",
            messages=[{"role":"user","content":prompt}])
        raw=response.choices[0].message.content.strip()
        with _progress_lock:
            progress_state["message"] = "Parsing GPT-5 response..."
            progress_state["log"].append("Parsing GPT-5 response...")
        match=re.search(r"\{.*\}|\[.*\]", raw, re.DOTALL)
        if not match: raise Exception("❌ No JSON object found in GPT output.")
        data=json.loads(match.group(0))

        # Normalize: if top-level is a list, wrap into dict with expected keys
        if isinstance(data, list):
            data = {"themes": {}, "gaps": data, "synthesis": {}, "references": []}

        themes=data.get("themes",{})
        gaps=data.get("gaps",[])
        synth=data.get("synthesis",{})
        refs=data.get("references",[])

        # Normalize synthesis values and provide fallbacks
        if isinstance(synth, dict):
            normalized = {}
            for k, v in list(synth.items()):
                if isinstance(v, dict):
                    v = " ".join(str(x) for x in v.values() if isinstance(x, (str, int, float)))
                if not isinstance(v, str):
                    v = str(v)
                normalized[str(k)] = v.strip()
            synth = normalized
        else:
            if isinstance(synth, str) and synth.strip():
                synth = {"Integrative": synth.strip()}
            else:
                synth = {}

        if not synth and ("Integrative" in data or "summary" in data):
            rv = data.get("Integrative", data.get("summary"))
            if isinstance(rv, dict):
                rv = " ".join(str(x) for x in rv.values() if isinstance(x, (str, int, float)))
            if not isinstance(rv, str):
                rv = str(rv)
            synth = {"Integrative": rv.strip()}

        with _progress_lock:
            progress_state["message"] = "Writing Excel..."
            progress_state["log"].append("Writing Excel...")
            progress_state["current"] = progress_state.get("total", 5) * 0.6

        # Type guards and small guard: continue even if themes empty
        if not isinstance(themes, dict):
            themes = {}
        if not isinstance(gaps, list):
            gaps = [str(gaps)] if gaps is not None else []
        if not isinstance(synth, dict):
            if isinstance(synth, str) and synth.strip():
                synth = {"Integrative": synth}
            else:
                synth = {}
        if not isinstance(refs, list):
            refs = [str(refs)] if refs is not None else []

        base=os.path.splitext(output_file)[0]
        excel_path=base+".xlsx"
        word_path=base+".docx"

        # ---------- Excel Export ----------
        with pd.ExcelWriter(excel_path,engine="xlsxwriter") as writer:
            rows=[]
            for th in themes_fixed:
                for art in themes.get(th,[]): rows.append([th,art])
            pd.DataFrame(rows,columns=["Theme","Article Title"]).to_excel(writer,sheet_name="Themes",index=False)
            pd.DataFrame(gaps,columns=["Gap"]).to_excel(writer,sheet_name="Gaps",index=False)
            syn_rows=[[k,(v.strip() if isinstance(v,str) else str(v))] for k,v in synth.items()]
            syn_rows+=[["Reference",r] for r in refs]
            pd.DataFrame(syn_rows,columns=["Theme","Content"]).to_excel(writer,sheet_name="Synthesis",index=False)

        # ---------- Word Export ----------
        with _progress_lock:
            progress_state["message"] = "Writing Word..."
            progress_state["log"].append("Writing Word...")
            progress_state["current"] = progress_state.get("total", 5) * 0.8
        doc=Document()
        doc.add_heading("Literature Synthesis Results",0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        for th in themes_fixed:
            text=synth.get(th,"").strip()
            if not text: continue
            doc.add_heading(th,level=2)
            doc.add_paragraph(text)
            doc.add_page_break()

        doc.add_heading("Research Gaps",level=2)
        for i,g in enumerate(gaps,1): doc.add_paragraph(f"{i}. {g}")
        doc.add_heading("References",level=2)
        for r in refs: doc.add_paragraph("• "+r)
        doc.save(word_path)

        with _progress_lock:
            progress_state.update({"done":True,"active":False,"message":
                f"✅ Synthesis complete. Saved {excel_path} and {word_path}.",
                "current":progress_state["total"]})
            progress_state["log"].append("Completed")
            progress_state["message"] = "Completed"
            progress_state["current"] = progress_state.get("total", 5)
            progress_state["done"] = True
            progress_state["active"] = False
    except Exception as e:
        with _progress_lock:
            progress_state.update({"done":True,"active":False,"message":f"❌ Error: {e}"})

# ---------- ROUTES ----------
@app.route("/",methods=["GET","POST"])
def index():
    if request.method=="POST":
        lit_file=request.form["lit_file"]
        output_file=request.form["output_file"]
        global cancel_event,progress_thread
        with _progress_lock:
            if progress_state.get("active") and not progress_state.get("done"):
                return jsonify({"ok":False,"error":"A run is already in progress"})
            progress_state.update({"active":True,"done":False,"message":"Starting...","log":[]})
        cancel_event=Event()
        progress_thread=Thread(target=run_synthesis,args=(lit_file,output_file),daemon=True)
        progress_thread.start()
        return jsonify({"ok":True,"started":True})
    return render_template_string(HTML)

@app.route("/progress")
def progress():
    with _progress_lock: return jsonify(progress_state)

@app.route("/cancel",methods=["POST"])
def cancel():
    global cancel_event
    cancel_event.set()
    with _progress_lock:
        progress_state.update({"active":False,"done":True,"message":"⚠️ Cancelled by user."})
    return jsonify({"ok":True,"canceled":True})

if __name__=="__main__":
    app.run(debug=True,threaded=True,port=5001)
