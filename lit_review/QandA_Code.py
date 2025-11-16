# ENTRYPOINT: flask_app
# PURPOSE: Launches a web interface for Question & Answer extraction and synthesis.
# HOW TO RUN:  python lit_review/QandA_Code.py

import os, json, pandas as pd, pdfplumber, docx
from flask import Flask, render_template_string, request, jsonify
from threading import Lock, Thread, Event
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from datetime import datetime
import xlsxwriter

# ---------- SETUP ----------
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
app = Flask(__name__)

_progress_lock = Lock()
cancel_event = Event()
progress_state = {"active": False, "done": False, "message": "", "log": [], "current": 0, "total": 0}
progress_thread = None

# ---------- FRONTEND ----------
HTML = """
<!DOCTYPE html>
<html>
<head>
<title>Question & Answer Tool</title>
<style>
body {font-family:Arial;margin:40px;}
label{display:block;margin-top:12px;font-weight:bold;}
input[type=text], textarea {width:80%;padding:8px;}
textarea {height:120px;}
button{margin-top:20px;padding:10px 20px;background:#0078D4;color:white;border:none;border-radius:5px;cursor:pointer;}
button:hover{background:#005ea0;}
.progress-wrap{margin-top:24px;max-width:800px;}
.bar{width:100%;background:#eee;border-radius:6px;overflow:hidden;height:16px;}
.bar>span{display:block;height:100%;background:#4caf50;width:0%;transition:width .25s ease;}
.log{margin-top:12px;background:#111;color:#0f0;padding:10px;border-radius:6px;height:180px;overflow:auto;
      font-family:Consolas,monospace;font-size:12px;}
.hide{display:none;}
.success{color:#0a7d00;font-weight:bold;}
.cancel-btn{margin-left:10px;padding:6px 12px;background:#c62828;color:#fff;border:none;border-radius:4px;cursor:pointer;}
.cancel-btn:hover{background:#a61e1e;}
</style>
</head>
<body>
<h2>dY's Question & Answer Tool</h2>

<form id="qaForm" method="POST" onsubmit="return false;">
  <label>Folder Path:</label>
  <input type="text" name="folder" required>
  <label>Research Question:</label>
  <textarea name="question" required></textarea>
  <label>Output Excel File (.xlsx):</label>
  <input type="text" name="output" required>
  <button type="submit">Run Q & A</button>
</form>

<div id="liveProgress" class="progress-wrap hide">
  <div class="bar"><span id="barFill"></span></div>
  <button id="cancelBtn" class="cancel-btn hide" type="button">Cancel</button>
  <div class="log" id="log"></div>
  <div id="finalMsg" class="success hide"></div>
</div>

<script>
window.addEventListener("DOMContentLoaded",()=>{
 const form=document.getElementById('qaForm');
 const live=document.getElementById('liveProgress');
 const bar=document.getElementById('barFill');
 const logBox=document.getElementById('log');
 const finalMsg=document.getElementById('finalMsg');
 const cancelBtn=document.getElementById('cancelBtn');
 let poller=null;

 function startPoll(){
   poller=setInterval(async()=>{
     const r=await fetch('/progress');const p=await r.json();
     if(p.total>0){bar.style.width=Math.min(100,Math.round(p.current/p.total*100))+'%';}
     if(p.log){logBox.textContent=p.log.join('\\n');logBox.scrollTop=logBox.scrollHeight;}
     if(p.done){clearInterval(poller);poller=null;cancelBtn.classList.add('hide');
       finalMsg.textContent=p.message;finalMsg.classList.remove('hide');}
   },1000);
 }

 form.addEventListener('submit',async e=>{
   e.preventDefault();e.stopPropagation();
   live.classList.remove('hide');cancelBtn.classList.remove('hide');
   finalMsg.classList.add('hide');logBox.textContent='';bar.style.width='0%';
   const fd=new FormData(form);
   await fetch('/',{method:'POST',body:fd});
   startPoll();
 });

 cancelBtn.onclick=async()=>{await fetch('/cancel',{method:'POST'});cancelBtn.disabled=true;};
});
</script>
</body>
</html>
"""

# ---------- HELPERS ----------
def extract_text_from_pdf(filepath):
    text=[]
    with pdfplumber.open(filepath) as pdf:
        for i,page in enumerate(pdf.pages,start=1):
            page_text=page.extract_text()
            if page_text:text.append(f"[Page {i}] {page_text}")
    return "\n".join(text)

def extract_text_from_docx(filepath):
    d=docx.Document(filepath)
    text=[f"[Paragraph {i+1}] {p.text.strip()}" for i,p in enumerate(d.paragraphs) if p.text.strip()]
    return "\n".join(text)

def extract_text_from_txt(filepath):
    with open(filepath,"r",encoding="utf-8") as f:
        lines=[f"[Line {i+1}] {line.strip()}" for i,line in enumerate(f.readlines()) if line.strip()]
    return "\n".join(lines)

def load_all_articles(folder):
    articles=[]
    for file in os.listdir(folder):
        fp=os.path.join(folder,file)
        if file.lower().endswith(".pdf"): text=extract_text_from_pdf(fp)
        elif file.lower().endswith(".docx"): text=extract_text_from_docx(fp)
        elif file.lower().endswith(".txt"): text=extract_text_from_txt(fp)
        else: continue
        if text: articles.append((file,text))
    return articles

# ---------- CORE ----------
def analyze_articles(question,articles):
    findings=[]
    total=len(articles)
    for i,(filename,text) in enumerate(articles,start=1):
        with _progress_lock:
            progress_state["message"]=f"Analyzing article {i} of {total}: {filename}"
            progress_state["log"].append(progress_state["message"])
            progress_state["current"]=i
        if cancel_event.is_set():return findings

        prompt=f"""
You are a research assistant. 
Question: {question}
Article: {filename}

Text:
{text}

Task:
- Identify any content that helps answer the question.
- Provide BOTH:
   1. A short paraphrased point (academic style) with in-text citation (file + page/paragraph/line).
   2. The raw excerpt from the text.

If no relevant information is found, return "No relevant info".
"""
        resp=client.chat.completions.create(model="gpt-5",
            messages=[{"role":"user","content":prompt}])
        ans=resp.choices[0].message.content.strip()
        if "No relevant info" not in ans:
            findings.append((filename,ans))
    return findings

def synthesize_answer(question,findings):
    combined="\n".join([f"- {f}: {p}" for f,p in findings])
    prompt=f"""
You are a research assistant. 
Write a synthesized academic-style answer with in-text citations.

Question: {question}

Findings:
{combined}
"""
    resp=client.chat.completions.create(model="gpt-5",
        messages=[{"role":"user","content":prompt}])
    return resp.choices[0].message.content.strip()

def run_qanda(folder,question,output):
    try:
        global cancel_event
        with _progress_lock:
            progress_state.update({"active":True,"done":False,"message":"Extracting full articles...","log":["Extracting full articles..."],"current":0})
        if cancel_event.is_set():return
        arts=load_all_articles(folder)
        total=len(arts)
        with _progress_lock:
            progress_state.update({"message":f"Scanning {total} articles with GPT-5...","log":["Scanning articles..."],"total":total})

        finds=analyze_articles(question,arts)
        with _progress_lock:
            progress_state["message"]="Synthesizing final answer..."
            progress_state["log"].append("Synthesizing final answer...")
            progress_state["current"]=total

        final=synthesize_answer(question,finds)

        # ---------- SAVE EXCEL ----------
        base=os.path.splitext(output)[0]
        excel_path=base+".xlsx"
        word_path=base+".docx"
        with pd.ExcelWriter(excel_path,engine="xlsxwriter") as writer:
            pd.DataFrame([[final]],columns=["Final Answer"]).to_excel(writer,sheet_name="Summary",index=False)
            pd.DataFrame(finds,columns=["File","Relevant Findings"]).to_excel(writer,sheet_name="Findings",index=False)

        # ---------- SAVE WORD ----------
        doc=Document()
        doc.add_heading("Question & Answer Summary",0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_heading("Question",level=2)
        doc.add_paragraph(question)
        doc.add_heading("Synthesized Answer",level=2)
        doc.add_paragraph(final)
        doc.add_page_break()
        doc.add_heading("Findings by Article",level=2)
        for f,p in finds:
            doc.add_heading(f,level=3)
            doc.add_paragraph(p)
        doc.save(word_path)

        with _progress_lock:
            progress_state.update({"done":True,"active":False,"message":f"✅ Q&A complete. Saved {excel_path} and {word_path}.","current":total})
            progress_state["log"].append("Completed.")
    except Exception as e:
        with _progress_lock:
            progress_state.update({"done":True,"active":False,"message":f"❌ Error: {e}"})

# ---------- ROUTES ----------
@app.route("/",methods=["GET","POST"])
def index():
    if request.method=="POST":
        folder=request.form["folder"];question=request.form["question"];output=request.form["output"]
        global cancel_event,progress_thread
        with _progress_lock:
            if progress_state.get("active") and not progress_state.get("done"):
                return jsonify({"ok":False,"error":"A run is already in progress"})
            progress_state.update({"active":True,"done":False,"message":"Starting...","log":[]})
        cancel_event=Event()
        progress_thread=Thread(target=run_qanda,args=(folder,question,output),daemon=True)
        progress_thread.start()
        return jsonify({"ok":True,"started":True})
    return render_template_string(HTML)

@app.route("/progress")
def progress(): 
    with _progress_lock: return jsonify(progress_state)

@app.route("/cancel",methods=["POST"])
def cancel():
    global cancel_event
    cancel_event.set()
    with _progress_lock:
        progress_state.update({"active":False,"done":True,"message":"⚠️ Cancelled by user."})
    return jsonify({"ok":True,"canceled":True})

if __name__=="__main__":
    app.run(debug=True,threaded=True,port=5002)
