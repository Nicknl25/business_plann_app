# ENTRYPOINT: flask_app
# PURPOSE: Launches a small web interface for Literature Review extraction.
# HOW TO RUN:  python lit_review/Lit_Review_Code.py

import os
import json
import pandas as pd
import pdfplumber
from PyPDF2 import PdfReader
from flask import Flask, render_template_string, request, jsonify
from openai import OpenAI
from dotenv import load_dotenv
import xlwings as xw
from threading import Lock, Thread, Event
from docx import Document  # ✅ Word export
from docx.shared import Pt

# Load environment variables
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = Flask(__name__)

# Global progress state
progress_state = {
    "active": False,
    "total": 0,
    "current": 0,
    "file": "",
    "log": [],
    "done": False,
    "error": None,
    "saved": 0,
    "output": "",
    "message": "",
}
_progress_lock = Lock()
cancel_event = Event()
progress_thread = None

# ---------------- HTML INTERFACE ----------------
HTML = """
<!DOCTYPE html>
<html>
<head>
<title>Literature Review Tool</title>
<style>
body { font-family: Arial, sans-serif; margin: 40px; }
label { display:block; margin-top:12px; font-weight:bold; }
input[type=text], textarea { width:80%; padding:8px; }
textarea { height:120px; }
button { margin-top:20px; padding:10px 20px; background:#0078D4; color:white;
  border:none; border-radius:5px; cursor:pointer;}
button:hover { background:#005ea0; }
.result { margin-top:25px; white-space:pre-wrap; background:#f8f8f8;
  padding:10px; border-radius:5px; }
.progress-wrap { margin-top:24px; max-width:800px; }
.bar { width:100%; background:#eee; border-radius:6px; overflow:hidden; height:16px; }
.bar > span { display:block; height:100%; background:#4caf50; width:0%; transition: width .25s ease; }
.metrics { margin-top:8px; color:#333; font-size:14px; }
.log { margin-top:12px; background:#111; color:#0f0; padding:10px; border-radius:6px;
  height:180px; overflow:auto; font-family:Consolas,monospace; font-size:12px; }
.hide { display:none; }
.success { color:#0a7d00; font-weight:bold; }
.cancel { color:#cc6600; font-weight:bold; }
.cancel-btn { margin-left:12px; padding:6px 12px; background:#c62828; color:#fff;
  border:none; border-radius:4px; cursor:pointer; }
.cancel-btn:hover { background:#a61e1e; }
</style>
</head>
<body>
<h2>dY's Literature Review Automation Tool</h2>

<form id="extractForm" method="POST">
  <label>Folder Path:</label>
  <input type="text" name="folder" required>
  <label>Topic / Research Focus:</label>
  <textarea name="topic" required></textarea>
  <label>Output Excel File (.xlsm):</label>
  <input type="text" name="output" required>
  <label>Sheet Name:</label>
  <input type="text" name="sheet" required>
  <button type="submit">Run Extraction</button>
</form>

{% if result %}
  <div class="result">{{ result }}</div>
{% endif %}

<div id="liveProgress" class="progress-wrap hide">
  <div class="bar"><span id="barFill" style="width:0%"></span></div>
  <button type="button" id="cancelBtn" class="cancel-btn" style="display:none;">Cancel</button>
  <div class="metrics" id="metrics">Idle</div>
  <div class="log" id="log"></div>
  <div class="result" id="finalMsg" style="display:none;"></div>
  <div class="result" id="errorMsg" style="display:none;color:#b00020;"></div>
  <div class="metrics" id="downloadHint" style="display:none;">Excel & Word exports have been saved to your specified path.</div>
</div>

<script>
window.addEventListener("DOMContentLoaded", () => {
  const form = document.getElementById('extractForm');
  const live = document.getElementById('liveProgress');
  const barFill = document.getElementById('barFill');
  const metrics = document.getElementById('metrics');
  const logBox = document.getElementById('log');
  const finalMsg = document.getElementById('finalMsg');
  const errorMsg = document.getElementById('errorMsg');
  const downloadHint = document.getElementById('downloadHint');
  const cancelBtn = document.getElementById('cancelBtn');
  let pollTimer = null;

  function resetUI(){
    live.classList.remove('hide');
    barFill.style.width = '0%';
    metrics.textContent = 'Starting...';
    logBox.textContent = '';
    finalMsg.style.display = 'none';
    errorMsg.style.display = 'none';
    downloadHint.style.display = 'none';
    document.querySelector('.bar').style.display = 'block';
    metrics.style.display = 'block';
    logBox.style.display = 'block';
    cancelBtn.style.display = 'inline-block';
    cancelBtn.disabled = false;
  }

  async function poll(){
    try {
      const r = await fetch('/progress', { cache: 'no-store' });
      if (!r.ok) return;
      const p = await r.json();
      const total = p.total || 0;
      const current = p.current || 0;
      const file = p.file || '';
      const done = !!p.done;
      const percent = total > 0 ? Math.min(100, Math.round((current / total) * 100)) : 0;
      barFill.style.width = percent + '%';
      if (p.active && !done && total > 0) {
        metrics.textContent = 'Processing file ' + Math.min(current + 1, total) + ' of ' + total + (file ? (': ' + file) : '');
      } else if (!done) {
        metrics.textContent = 'Preparing...';
      }
      if (Array.isArray(p.log)) {
        logBox.textContent = p.log.join('\\n');
        logBox.scrollTop = logBox.scrollHeight;
      }
      cancelBtn.style.display = (p.active && !done) ? 'inline-block' : 'none';
      if (done) {
        stopPolling();
        document.querySelector('.bar').style.display = 'none';
        metrics.style.display = 'none';
        logBox.style.display = 'none';
        cancelBtn.style.display = 'none';
        const msg = p.message || ('✅ Completed — saved ' + (p.saved || 0) + ' entries to Excel & Word');
        finalMsg.textContent = msg;
        finalMsg.className = msg.includes('Cancelled') ? 'result cancel' : 'result success';
        finalMsg.style.display = 'block';
        if (!msg.includes('Cancelled')) downloadHint.style.display = 'block';
      }
    } catch(e) {}
  }

  function startPolling(){ stopPolling(); pollTimer = setInterval(poll, 1000); poll(); }
  function stopPolling(){ if(pollTimer){clearInterval(pollTimer);pollTimer=null;} }

  form.addEventListener('submit', async e=>{
    e.preventDefault();
    resetUI();
    startPolling();
    try {
      const r = await fetch('/', { method:'POST', body:new FormData(form), headers:{'X-Requested-With':'XMLHttpRequest'} });
      const data = await r.json();
      if(!(data && data.ok)){
        errorMsg.textContent = (data && data.error) ? data.error : 'Unexpected error starting job';
        errorMsg.style.display = 'block';
      }
    } catch(err){
      errorMsg.textContent = 'Request failed: ' + err;
      errorMsg.style.display = 'block';
    }
  });

  cancelBtn.addEventListener('click', async ()=>{
    cancelBtn.disabled = true;
    try { await fetch('/cancel', { method:'POST' }); } catch(e) {}
    stopPolling();
    document.querySelector('.bar').style.display = 'none';
    metrics.style.display = 'none';
    logBox.style.display = 'none';
    cancelBtn.style.display = 'none';
    finalMsg.textContent = '⚠️ Cancelled by user.';
    finalMsg.className = 'result cancel';
    finalMsg.style.display = 'block';
    downloadHint.style.display = 'none';
  });
});
</script>
</body>
</html>
"""
# -----------------------------------------------------------


def extract_text(path):
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for p in pdf.pages:
                text += p.extract_text() or ""
    except Exception:
        reader = PdfReader(path)
        for p in reader.pages:
            text += p.extract_text() or ""
    return text


def analyze(text, topic):
    prompt = f"""
    You are an academic assistant. Analyze this paper and return ONLY valid JSON.

    {{
      "Title": "string",
      "Author": "string",
      "Publication Date": "string",
      "APA Citation": "string",
      "Overall": "string",
      "Methods": "string",
      "Conclusion": "string",
      "Score": "integer"
    }}

    Topic: "{topic}"
    Document:
    {text[:6000]}
    """
    resp = client.chat.completions.create(
        model="gpt-5",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(resp.choices[0].message.content)


def run_extraction(folder, topic, output, sheet):
    try:
        global cancel_event
        if cancel_event.is_set():
            with _progress_lock:
                progress_state.update({"active": False, "done": True, "message": "⚠️ Cancelled by user."})
            return

        files = [n for n in os.listdir(folder) if n.lower().endswith('.pdf')]
        files.sort()
        results = []
        with _progress_lock:
            progress_state.update({
                "active": True, "done": False, "error": None, "total": len(files),
                "current": 0, "file": "", "log": [], "saved": 0, "output": output, "message": ""
            })

        for idx, name in enumerate(files):
            if cancel_event.is_set():
                with _progress_lock:
                    progress_state.update({"active": False, "done": True, "message": "⚠️ Cancelled by user."})
                return

            path = os.path.join(folder, name)
            with _progress_lock:
                progress_state["file"] = name
                progress_state["log"].append(f"Processing file {idx+1} of {len(files)}: {name}")
                progress_state["current"] = idx

            text = extract_text(path)
            if not text.strip():
                with _progress_lock:
                    progress_state["log"].append(f"Skipped (no text): {name}")
                    progress_state["current"] = idx + 1
                continue

            try:
                parsed = analyze(text, topic)
                parsed["File"] = name
                parsed["Source"] = path
                results.append(parsed)
                with _progress_lock:
                    progress_state["log"].append(f"Completed: {name}")
                    progress_state["current"] = idx + 1
            except Exception as e:
                err_msg = f"Error: {name}: {e}"
                print(err_msg)
                with _progress_lock:
                    progress_state["log"].append(err_msg)
                    progress_state["current"] = idx + 1

        if not results:
            with _progress_lock:
                progress_state.update({"done": True, "active": False, "message": "No results found"})
            return

        if cancel_event.is_set():
            with _progress_lock:
                progress_state.update({"active": False, "done": True, "message": "⚠️ Cancelled by user."})
            return

        # ---------- Excel Export ----------
        df = pd.DataFrame(results)
        wb = xw.Book(output)
        if sheet in [s.name for s in wb.sheets]:
            sht = wb.sheets[sheet]
        else:
            sht = wb.sheets.add(sheet)
            sht.range("A1").value = list(df.columns)
        sht.range("A2").value = df.values.tolist()
        wb.save(); wb.close()

        # ---------- Word Export ----------
        try:
            docx_path = os.path.splitext(output)[0] + ".docx"
            document = Document()
            document.add_heading('Literature Review Results', 0)

            for _, row in df.iterrows():
                document.add_paragraph('----------------------------------------')

                # Title as Heading 2
                title = row.get('Title', '')
                if title:
                    document.add_heading(title, level=2)
                else:
                    document.add_heading("Untitled Document", level=2)

                document.add_paragraph(f"Author: {row.get('Author', '')}")
                document.add_paragraph(f"Publication Date: {row.get('Publication Date', '')}")
                document.add_paragraph(f"APA Citation: {row.get('APA Citation', '')}")
                document.add_paragraph(f"Score: {row.get('Score', '')}")

                document.add_paragraph("Overall:")
                document.add_paragraph(row.get("Overall", ""), style="List Bullet")

                document.add_paragraph("Methods:")
                document.add_paragraph(row.get("Methods", ""), style="List Bullet")

                document.add_paragraph("Conclusion:")
                document.add_paragraph(row.get("Conclusion", ""), style="List Bullet")

                document.add_paragraph(f"Source File: {row.get('Source', '')}")
                document.add_paragraph(f"File Name: {row.get('File', '')}")

                # Add page break after each article
                document.add_page_break()

            document.save(docx_path)
        except Exception as e:
            print("⚠️ Word export failed:", e)

        # ---------- Completion Update ----------
        with _progress_lock:
            progress_state.update({
                "done": True, "active": False, "saved": len(df),
                "message": f"Done - saved {len(df)} entries to {output} and Word document.",
                "current": progress_state.get("total", 0)
            })
    except Exception as e:
        with _progress_lock:
            progress_state.update({"done": True, "active": False, "error": str(e)})


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        folder = request.form["folder"]
        topic = request.form["topic"]
        output = request.form["output"]
        sheet = request.form["sheet"]

        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            global cancel_event, progress_thread
            with _progress_lock:
                if progress_state.get("active") and not progress_state.get("done"):
                    return jsonify({"ok": False, "error": "A run is already in progress"})
                progress_state.update({"active": True, "done": False, "error": None, "saved": 0, "log": []})
            cancel_event = Event()
            progress_thread = Thread(target=run_extraction, args=(folder, topic, output, sheet), daemon=True)
            progress_thread.start()
            return jsonify({"ok": True, "started": True})

        run_extraction(folder, topic, output, sheet)
        with _progress_lock:
            msg = progress_state.get("message") or ("Error: " + str(progress_state.get("error")))
        return render_template_string(HTML, result=msg)

    return render_template_string(HTML)


@app.route("/progress", methods=["GET"])
def progress():
    with _progress_lock:
        return jsonify(progress_state)


@app.route("/cancel", methods=["POST"])
def cancel():
    global cancel_event, progress_thread
    with _progress_lock:
        cancel_event.set()
        progress_state.update({"active": False, "done": True, "message": "⚠️ Cancelled by user."})
    try:
        if progress_thread and progress_thread.is_alive():
            progress_thread.join(timeout=0.1)
    except Exception:
        pass
    return jsonify({"ok": True, "canceled": True})


if __name__ == "__main__":
    app.run(debug=True, threaded=True)
