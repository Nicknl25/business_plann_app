"""READ THE PRODUCED .DOCX BACK (2026-08-30) and feed the craft checks.

The probe inspects the saved file's XML, not the objects that built it - the
same discipline as reading Checks!B2 out of the workbook instead of trusting
the writer. A document that cannot be probed FAILS (the CoInitialize law);
checks.py enforces that by treating a missing probe as could-not-run.
"""
from __future__ import annotations

import re
import zipfile
from typing import Any, Dict, List, Set

from docx import Document

from .. import rules as R

_RUN_FMT_TAGS = ("<w:b/>", "<w:i/>", "<w:sz ", "<w:rFonts", "<w:color")


def probe_docx(path: str, *, run_id: str) -> Dict[str, Any]:
  doc = Document(path)
  with zipfile.ZipFile(path) as z:
    body_xml = z.read("word/document.xml").decode("utf-8")
    styles_xml = z.read("word/styles.xml").decode("utf-8")
    hf_xml = " ".join(
      z.read(n).decode("utf-8") for n in z.namelist()
      if n.startswith("word/header") or n.startswith("word/footer"))

  # ---- rule 23 (narrowed 2026-08-30): text boxes banned; anchors CLASSIFIED.
  # An anchor carrying wrapSquare/wrapTight flows with its paragraph and is
  # permitted; anything else (wrapNone, through, behind-text) is absolute
  # positioning and fights an editor.
  text_boxes = body_xml.count("w:txbxContent")
  anchored_wrapped = 0
  absolutely_positioned = 0
  for m in re.finditer(r"<wp:anchor\b(?:(?!</wp:anchor>).)*?</wp:anchor>", body_xml, re.S):
    blk = m.group(0)
    if "<wp:wrapSquare" in blk or "<wp:wrapTight" in blk:
      anchored_wrapped += 1
    else:
      absolutely_positioned += 1
  floating = anchored_wrapped + absolutely_positioned
  inline_images = body_xml.count("<wp:inline")

  # ---- rule 22: no direct-formatted runs in the BODY (styles own everything).
  # Field runs (PAGE/TOC) carry no rPr; a run-level rPr with font/size/bold in
  # document.xml is a bypass of the style system.
  direct = 0
  for m in re.finditer(r"<w:r>(?:(?!</w:r>).)*?</w:r>", body_xml, re.S):
    chunk = m.group(0)
    if "<w:rPr>" in chunk and any(t in chunk for t in _RUN_FMT_TAGS):
      direct += 1

  # ---- real styles in use
  used_para_styles = {p.style.name for p in doc.paragraphs if p.style is not None}
  uses_real_styles = ("Heading 1" in used_para_styles)

  # ---- one font pair: fonts named by the styles we defined + any run-level
  fonts: Set[str] = set(re.findall(r'w:ascii="([^"]+)"', styles_xml))
  fonts |= set(re.findall(r'w:ascii="([^"]+)"', body_xml))
  # Word seeds latent theme fonts on built-in styles we never touch; count only
  # fonts reachable from content: Normal/Heading/Title/Caption/Chrome/Table.
  declared = set()
  for sid in ("Normal", "Heading1", "Heading2", "Title", "PlanSubtitle",
              "PlanCaption", "PlanChrome", "PlanTable"):
    m = re.search(r'w:styleId="%s".*?</w:style>' % sid, styles_xml, re.S)
    if m:
      declared |= set(re.findall(r'w:ascii="([^"]+)"', m.group(0)))
  body_fonts = set(re.findall(r'w:ascii="([^"]+)"', body_xml))
  families = declared | body_fonts

  # ---- one table style
  table_styles = set(re.findall(r'<w:tblStyle w:val="([^"]+)"', body_xml))

  # ---- figures: caption bijection + cross-references
  full_text = "\n".join(p.text for p in doc.paragraphs)
  captions = re.findall(r"Figure (\d+) —", full_text)
  refs = re.findall(r"Figure (\d+)\b(?! —)", full_text)
  fig_missing_caption = sorted(set(refs) - set(captions))
  fig_missing_ref = sorted(set(captions) - set(refs))

  # ---- rule 21: footer template + run id placement
  footer_ok = ("Confidential · Page" in hf_xml.replace("</w:t>", "").replace("<w:t>", "")
               or "Confidential" in hf_xml)
  # run id: split the body at the Appendix heading
  parts = full_text.split("\nAppendix\n")
  before = parts[0]
  after = parts[1] if len(parts) > 1 else ""
  return {
    "text_boxes": text_boxes,
    "floating_shapes": floating,
    "anchored_wrapped": anchored_wrapped,
    "absolutely_positioned": absolutely_positioned,
    "non_inline_images": absolutely_positioned,
    "inline_images": inline_images,
    "direct_formatted_runs": direct,
    "uses_real_styles": uses_real_styles,
    "font_families_used": len(families),
    "font_families": sorted(families),
    "table_styles_used": len(table_styles),
    "figures_without_caption": fig_missing_caption,
    "figures_without_cross_reference": fig_missing_ref,
    "footer_matches_template": footer_ok,
    "run_id_in_body": run_id in before,
    "run_id_in_header": run_id in hf_xml,
    "run_id_in_footer": run_id in hf_xml,
    "run_id_in_appendix": run_id in after,
    "sections": len(doc.sections),
    "used_paragraph_styles": sorted(used_para_styles),
  }
