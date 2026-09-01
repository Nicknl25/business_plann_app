"""THE FOUR BODY TABLES (Nick's ruling, 2026-08-31): condensed annual
statements, the assumptions table, sources & uses, annual debt amortization.
Everything else is appendix - a body full of tables reads like an appendix
that escaped.

Each builder returns a spec {title, columns, rows, note} or None when the
table cannot be built - same ABSENT discipline as the facts: a table with
nothing honest to say is omitted, never padded. write_table() renders a spec
with the shell's single "Plan Table" style: numbers right-aligned, no direct
run formatting.

Sources & uses balances BY CONSTRUCTION: both sides are read off the same
cash identity the model already balances, so a mismatch here means the model
is broken, not the table.
"""
from __future__ import annotations

import json
import math
from typing import Any, Dict, List, Optional, Sequence


def _j(v: Any) -> Any:
  if isinstance(v, (dict, list)):
    return v
  try:
    return json.loads(v) if v else {}
  except Exception:
    return {}


def _f(v: Any) -> Optional[float]:
  try:
    x = float(v)
    return x if math.isfinite(x) else None
  except (TypeError, ValueError):
    return None


def _quarters(finmo: Dict[str, Any]) -> Dict[int, Dict[str, Any]]:
  out: Dict[int, Dict[str, Any]] = {}
  for r in finmo.get("quarter_rows") or []:
    qi = _f(r.get("quarter_index"))
    if qi is not None:
      out[int(qi)] = r
  return out


def _ysum(q, key, year) -> Optional[float]:
  vals = [_f(q.get(i, {}).get(key)) for i in range(4 * year - 3, 4 * year + 1)]
  return None if any(v is None for v in vals) else float(sum(vals))


def _money(v: Optional[float]) -> str:
  if v is None or abs(v) < 0.5:
    return "—"
  return ("(%s)" % "{:,.0f}".format(abs(v))) if v < 0 else "{:,.0f}".format(v)


def _pct(v: Optional[float]) -> str:
  return "—" if v is None else "%.1f%%" % (v * 100)


# ---------------------------------------------------------------------------
# 1. CONDENSED ANNUAL STATEMENTS
# ---------------------------------------------------------------------------
_PL_LINES = (
  ("Revenue", "revenue", False),
  ("Cost of goods sold", "cogs", True),
  ("Gross profit", None, False),          # computed
  ("Payroll", "payroll", True),
  ("Marketing", "marketing", True),
  ("Lease / rent", "lease_rent", True),
  ("General & administrative", "g_and_a", True),
  ("EBITDA", "ebitda", False),
  ("Depreciation", "depreciation", True),
  ("Interest", "interest", True),
  ("Taxes", "taxes", True),
  ("Net income", "net_income", False),
)


def build_condensed_statements(draft: Dict[str, Any]) -> Optional[Dict[str, Any]]:
  q = _quarters(_j(draft.get("finmo_json")))
  if len([i for i in q if 1 <= i <= 20]) < 20:
    return None
  rows: List[List[str]] = []
  ysums = {key: [_ysum(q, key, y) for y in range(1, 6)]
           for _, key, _neg in _PL_LINES if key}
  if not all(v is not None for v in ysums["revenue"]):
    return None
  for label, key, neg in _PL_LINES:
    if key is None:   # gross profit
      vals = [ysums["revenue"][i] - (ysums["cogs"][i] or 0) for i in range(5)]
    else:
      vals = [(-(v or 0) if neg and v is not None else v) for v in ysums[key]]
      vals = [v if ysums[key][i] is not None else None for i, v in enumerate(vals)]
    rows.append([label] + [_money(v) for v in vals])
  return {"key": "condensed_statements", "title": "Five-Year Summary of Operations",
          "columns": ["", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
          "rows": rows,
          "emphasis_rows": [i for i, (lab, _k, _n) in enumerate(_PL_LINES)
                            if lab in ("Revenue", "Gross profit", "EBITDA", "Net income")],
          "note": "Costs shown in parentheses. Full quarterly statements are in the accompanying workbook."}


# ---------------------------------------------------------------------------
# 2. THE ASSUMPTIONS TABLE (driver -> value -> basis)
# ---------------------------------------------------------------------------
def build_assumptions_table(draft: Dict[str, Any]) -> Optional[Dict[str, Any]]:
  mi = _j(draft.get("model_input_json"))
  fin = _j(draft.get("financials_json"))
  ph = _j(draft.get("payroll_headcount"))
  si = (mi.get("solver_input") or {}) if isinstance(mi, dict) else {}
  q = _quarters(_j(draft.get("finmo_json")))
  rows: List[List[str]] = []

  jg = si.get("judged_growth") or {}
  if _f(jg.get("qoq_start")) is not None and _f(jg.get("qoq_end")) is not None:
    rows.append(["Quarterly revenue growth",
                 "%s easing to %s" % (_pct(_f(jg["qoq_start"])), _pct(_f(jg["qoq_end"]))),
                 "judged in the coherence review"])
  fb = si.get("fitted_bands") or {}
  cg = fb.get("cogs_percent_of_revenue") or {}
  c1, c20 = _f(cg.get("1")), _f(cg.get("20"))
  if c1 is not None and c20 is not None:
    rows.append(["Cost of goods sold", "%s of revenue, easing to %s" % (_pct(c1), _pct(c20)),
                 "fitted band grounded in the stated baseline"])
  elif _f(fin.get("cogs_percent_of_revenue")) is not None:
    rows.append(["Cost of goods sold", "%s of revenue" % _pct(_f(fin["cogs_percent_of_revenue"])),
                 "stated baseline"])
  if _f(fin.get("marketing_percent_of_revenue")) is not None:
    rows.append(["Marketing spend", "%s of revenue" % _pct(_f(fin["marketing_percent_of_revenue"])),
                 "stated baseline"])
  wc = (si.get("wc_judgment") or {}).get("drivers") or {}
  _WC = (("ar_days", "Receivable days"), ("ap_days", "Payable days"),
         ("inventory_days", "Inventory days"))
  for key, label in _WC:
    d = wc.get(key) or {}
    if d.get("applicable") and _f(d.get("q1")) is not None and _f(d.get("q20")) is not None:
      rows.append([label, "%.0f days moving to %.0f" % (float(d["q1"]), float(d["q20"])),
                   "judged from the operating terms described at intake"])
  cj = si.get("cash_judgment") or {}
  if _f(cj.get("buffer_months")) is not None:
    rows.append(["Minimum cash buffer", "%.1f months of operating cost" % float(cj["buffer_months"]),
                 "judged cash posture"])
  mb = si.get("margin_band_judgment") or {}
  t20 = _f((mb.get("q20") or {}).get("target"))
  if t20 is not None:
    rows.append(["Mature net-income margin", _pct(t20), "judged band, target at Year 5"])
  if 1 in q and _f(q[1].get("debt_interest_rate")) is not None and _f(q[1].get("debt_closing_balance")):
    rows.append(["Interest rate on debt", _pct(float(q[1]["debt_interest_rate"]) * 4.0),
                 "annualised from the model's quarterly rate"])
  winf = None
  for r in (ph.get("rows") or []):
    winf = _f(r.get("annual_wage_inflation_rate"))
    if winf is not None:
      break
  if winf is not None:
    rows.append(["Annual wage inflation", _pct(winf), "applied to every scheduled wage"])
  for r0 in ((mi.get("sections") or {}).get("expenses") or []):
    if str(r0.get("label")) == "Taxes":
      vals = r0.get("values") or []
      tv = _f(vals[1] if len(vals) > 1 else None)
      if tv is not None:
        rows.append(["Effective tax rate", _pct(tv), "the model's stated rate"])
      break
  if len(rows) < 3:
    return None
  return {"key": "assumptions", "title": "Key Assumptions",
          "columns": ["Driver", "Value", "Basis"], "rows": rows, "emphasis_rows": [],
          "note": "Each assumption is held in the financial model; the workbook's Model Inputs sheet carries the full set."}


# ---------------------------------------------------------------------------
# 3. SOURCES & USES (Year 1, off the cash identity)
# ---------------------------------------------------------------------------
def build_sources_and_uses(draft: Dict[str, Any]) -> Optional[Dict[str, Any]]:
  q = _quarters(_j(draft.get("finmo_json")))
  if 0 not in q or len([i for i in q if 1 <= i <= 4]) < 4:
    return None
  Y = lambda k: _ysum(q, k, 1) or 0.0
  opening_cash = _f(q[0].get("cash")) or 0.0
  ocf = Y("operating_cash_flow")
  debt_in = Y("debt_issuance")
  capex = Y("capital_expenditures")
  debt_out = Y("debt_repayment")
  lease_out = Y("lease_principal_repayments")
  dists = Y("distributions") or Y("owner_distributions")
  # equity IN is the residual of the model's own financing line - never a sum
  # of balance columns. `other_equity` is a BALANCE repeated per quarter; the
  # first build summed it as a flow and the 2% guard refused Northgate's
  # table for a $10M phantom source (2026-09-01). The guard earned its keep.
  equity_in = Y("financing_cash_flow") - (debt_in - debt_out - lease_out - dists)
  closing_cash = _f(q[4].get("cash")) or 0.0

  sources = [("Cash on hand at the start", opening_cash),
             ("Debt drawn in Year 1", debt_in)]
  if equity_in > 0.5:
    sources.append(("Equity contributed", equity_in))
  if ocf > 0:
    sources.append(("Cash generated by operations", ocf))
  uses = [("Capital expenditure", capex),
          ("Debt principal repaid", debt_out)]
  if lease_out:
    uses.append(("Lease principal repaid", lease_out))
  if dists:
    uses.append(("Owner distributions", dists))
  if equity_in < -0.5:
    uses.append(("Other financing outflows", -equity_in))
  if ocf < 0:
    uses.append(("Cash absorbed by operations", -ocf))
  uses.append(("Cash on hand at the end of Year 1", closing_cash))

  st, ut = sum(v for _, v in sources), sum(v for _, v in uses)
  # the identity: both sides come off the model's balanced cash flow. Beyond
  # rounding-scale drift the model itself is broken - refuse the table.
  if st <= 0 or abs(st - ut) / max(st, 1.0) > 0.02:
    return None
  rows = [[label, _money(v)] for label, v in sources]
  rows.append(["Total sources", _money(st)])
  n_src = len(rows)
  rows += [[label, _money(v)] for label, v in uses]
  rows.append(["Total uses", _money(ut)])
  return {"key": "sources_and_uses", "title": "Sources and Uses of Capital (Year 1)",
          "columns": ["", "Amount"], "rows": rows,
          "emphasis_rows": [n_src - 1, len(rows) - 1],
          "note": "Both sides are read off the model's balanced cash flow; they agree because the model does."}


# ---------------------------------------------------------------------------
# 4. ANNUAL DEBT AMORTIZATION
# ---------------------------------------------------------------------------
def build_debt_amortization(draft: Dict[str, Any]) -> Optional[Dict[str, Any]]:
  q = _quarters(_j(draft.get("finmo_json")))
  if len([i for i in q if 1 <= i <= 20]) < 20:
    return None
  rows: List[List[str]] = []
  any_debt = False
  for y in range(1, 6):
    opening = _f(q[4 * y - 3].get("debt_opening_balance"))
    drawn = _ysum(q, "debt_issuance", y)
    principal = _ysum(q, "debt_repayment", y)
    interest = _ysum(q, "debt_interest_expense_only", y)
    if interest is None:
      interest = _ysum(q, "debt_interest_expense", y)
    closing = _f(q[4 * y].get("debt_closing_balance"))
    if opening is None or closing is None:
      return None
    if any(v and abs(v) > 0.5 for v in (opening, drawn, principal, closing)):
      any_debt = True
    rows.append(["Year %d" % y, _money(opening), _money(drawn or 0),
                 _money(principal or 0), _money(interest or 0), _money(closing)])
  if not any_debt:
    return None   # no debt anywhere in the plan - nothing to amortize
  return {"key": "debt_amortization", "title": "Debt Amortization by Year",
          "columns": ["", "Opening", "Drawn", "Principal paid", "Interest", "Closing"],
          "rows": rows, "emphasis_rows": [],
          "note": "The quarterly Debt Schedule in the workbook carries the payment-level detail."}


# ---------------------------------------------------------------------------
# 5. ANNUAL BALANCE SHEET (year-end balances; ppe is carried NET in the model)
# ---------------------------------------------------------------------------
_BS_LINES = (
  ("Cash", "cash", "a"), ("Accounts receivable", "accounts_receivable", "a"),
  ("Inventory", "inventory", "a"), ("Prepaid expenses", "prepaid_expenses", "a"),
  ("Total current assets", "current_assets", "T"),
  ("Property & equipment, net", "ppe", "a"), ("Right-of-use asset", "right_of_use_asset", "a"),
  ("Total assets", "total_assets", "T"),
  ("Accounts payable", "accounts_payable", "l"), ("Deferred revenue", "deferred_revenue", "l"),
  ("Short-term debt", "short_term_debt", "l"),
  ("Total current liabilities", "current_liabilities", "T"),
  ("Long-term debt", "long_term_debt", "l"), ("Lease obligation", "capital_lease_obligation", "l"),
  ("Total liabilities", "total_liabilities", "T"),
  ("Owner's capital", "owners_capital", "e"), ("Retained earnings", "retained_earnings", "e"),
  ("Total equity", "total_equity", "T"),
  ("Total liabilities & equity", "total_liabilities_and_equity", "T"),
)


def build_balance_sheet(draft: Dict[str, Any]) -> Optional[Dict[str, Any]]:
  q = _quarters(_j(draft.get("finmo_json")))
  if len([i for i in q if 1 <= i <= 20]) < 20:
    return None
  ends = [q[4 * y] for y in range(1, 6)]
  # the identity, per year: refuse rather than print a sheet that doesn't balance
  for r in ends:
    ta, le = _f(r.get("total_assets")), _f(r.get("total_liabilities_and_equity"))
    if ta is None or le is None or abs(ta - le) > max(1.0, 0.005 * abs(ta)):
      return None
  rows, emph = [], []
  for label, key, kind in _BS_LINES:
    vals = [_f(r.get(key)) for r in ends]
    if kind != "T" and not any(v and abs(v) > 0.5 for v in vals):
      continue   # an all-zero detail line says nothing; totals always print
    if kind == "T":
      emph.append(len(rows))
    rows.append([label] + [_money(v) for v in vals])
  return {"key": "balance_sheet", "title": "Balance Sheet at Each Year End",
          "columns": ["", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
          "rows": rows, "emphasis_rows": emph,
          "note": "Year-end balances; the quarterly balance sheet is in the accompanying workbook."}


# ---------------------------------------------------------------------------
# 6. ANNUAL CASH FLOW (flows summed; subtotals are the model's own)
# ---------------------------------------------------------------------------
def build_cash_flow(draft: Dict[str, Any]) -> Optional[Dict[str, Any]]:
  q = _quarters(_j(draft.get("finmo_json")))
  if len([i for i in q if 1 <= i <= 20]) < 20:
    return None
  Y = lambda k, y: _ysum(q, k, y) or 0.0
  rows, emph = [], []
  def line(label, vals, total=False, neg=False, skip_zero=True):
    if skip_zero and not any(abs(v) > 0.5 for v in vals):
      return
    if total:
      emph.append(len(rows))
    rows.append([label] + [_money(-v if neg else v) for v in vals])
  ni = [Y("net_income", y) for y in range(1, 6)]
  dep = [Y("depreciation", y) for y in range(1, 6)]
  dca = [Y("changes_in_current_assets", y) for y in range(1, 6)]
  dcl = [Y("changes_in_current_liabilities", y) for y in range(1, 6)]
  ocf = [Y("operating_cash_flow", y) for y in range(1, 6)]
  other_op = [ocf[i] - (ni[i] + dep[i] + dca[i] + dcl[i]) for i in range(5)]
  line("Net income", ni, skip_zero=False)
  line("Depreciation", dep)
  line("Change in current assets", dca)
  line("Change in current liabilities", dcl)
  line("Other operating items", other_op)
  line("Cash from operations", ocf, total=True, skip_zero=False)
  capex = [Y("capital_expenditures", y) for y in range(1, 6)]
  icf = [Y("investing_cash_flow", y) for y in range(1, 6)]
  line("Capital expenditure", capex, neg=True)
  other_inv = [icf[i] + capex[i] for i in range(5)]
  line("Other investing items", other_inv)
  line("Cash from investing", icf, total=True, skip_zero=False)
  d_in = [Y("debt_issuance", y) for y in range(1, 6)]
  d_out = [Y("debt_repayment", y) for y in range(1, 6)]
  l_out = [Y("lease_principal_repayments", y) for y in range(1, 6)]
  dist = [Y("distributions", y) or Y("owner_distributions", y) for y in range(1, 6)]
  fcf = [Y("financing_cash_flow", y) for y in range(1, 6)]
  eq = [fcf[i] - (d_in[i] - d_out[i] - l_out[i] - dist[i]) for i in range(5)]
  line("Debt drawn", d_in)
  line("Debt repaid", d_out, neg=True)
  line("Lease principal repaid", l_out, neg=True)
  line("Equity contributed (net of other financing)", eq)
  line("Owner distributions", dist, neg=True)
  line("Cash from financing", fcf, total=True, skip_zero=False)
  net = [ocf[i] + icf[i] + fcf[i] for i in range(5)]
  line("Net change in cash", net, total=True, skip_zero=False)
  ending = [_f(q[4 * y].get("cash")) or 0.0 for y in range(1, 6)]
  line("Cash at year end", ending, total=True, skip_zero=False)
  # identity: opening + net == ending each year, else refuse
  opening = _f(q[0].get("cash")) or 0.0
  for i in range(5):
    if abs(opening + net[i] - ending[i]) > max(1.0, 0.005 * abs(ending[i])):
      return None
    opening = ending[i]
  return {"key": "cash_flow", "title": "Cash Flow by Year",
          "columns": ["", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"],
          "rows": rows, "emphasis_rows": emph,
          "note": "Outflows in parentheses. Subtotals are the model's own; component lines reconcile to them."}


BODY_TABLE_BUILDERS = (build_condensed_statements, build_balance_sheet, build_cash_flow,
                       build_assumptions_table, build_sources_and_uses, build_debt_amortization)


def build_body_tables(draft: Dict[str, Any]) -> List[Dict[str, Any]]:
  out = []
  for fn in BODY_TABLE_BUILDERS:
    try:
      spec = fn(draft)
    except Exception:
      spec = None
    if spec:
      out.append(spec)
  return out


# ---------------------------------------------------------------------------
# THE DOCX WRITER - one style, numbers right-aligned, no direct formatting
# ---------------------------------------------------------------------------
def write_table(doc, spec: Dict[str, Any]):
  """Render a spec with the shell's single 'Plan Table' style. Numeric cells
  (everything after the first column) right-align via paragraph alignment -
  never via run formatting."""
  from docx.enum.text import WD_ALIGN_PARAGRAPH
  cols = spec["columns"]
  t = doc.add_table(rows=1 + len(spec["rows"]), cols=len(cols))
  t.style = doc.styles["Plan Table"]
  for j, c in enumerate(cols):
    cell = t.cell(0, j)
    cell.text = str(c)
    if j > 0:
      cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
  for i, row in enumerate(spec["rows"]):
    for j, v in enumerate(row):
      cell = t.cell(i + 1, j)
      cell.text = str(v)
      if j > 0 and spec["key"] != "assumptions":
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
  if spec.get("note"):
    doc.add_paragraph(spec["note"], style="Chart Note") if "Chart Note" in [s.name for s in doc.styles] \
      else doc.add_paragraph(spec["note"])
  return t
