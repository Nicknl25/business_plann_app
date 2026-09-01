"""SECTION-DRAFT ASSEMBLER (2026-09-01) - authored prose into the docx shell.

Nick: "Section output goes to C:\\dev\\Client Written Plans, not buried in
_writing_business/. Name them so I can tell what I'm opening." So every
authored section lands as a real document - the shell's styles, running
header, stamped footer - named
    "<Business Name> -- <Section Title> -- MM-DD-YYYY HH-MM-SS.docx"
(the workbook convention with the section named). The title page says
SECTION DRAFT so nobody mistakes one for a delivered plan; the run identifier
sits in a minimal appendix, never on a client-facing page (rule 21).

Everything is a real Word style, including the note references: a character
style "Plan Note Ref" carries the superscript, so no run is ever directly
formatted (rule 22) and the probe stays clean.
"""
from __future__ import annotations

import datetime as _dt
import os
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE

from .. import rules as R
from ..checks import FACT_TOKEN, _SUPERSCRIPT_MARKER
from ..facts.catalog import FactCatalog
from . import renderer as REN

_UNSAFE = re.compile(r'[\\/:*?"<>|]+')


def _note_ref_style(doc: Document) -> None:
  st = doc.styles
  ref = st.add_style("Plan Note Ref", WD_STYLE_TYPE.CHARACTER)
  ref.font.superscript = True


def _runs(text: str, cat: FactCatalog) -> List[Tuple[str, str]]:
  """Split a sentence into (kind, value) runs: 'text' with tokens rendered
  through the one formatter, and 'noteref' for each [^n] marker."""
  def _sub(m):
    f = cat.get_quiet(m.group(1))
    return f.render() if f is not None else m.group(0)

  rendered = FACT_TOKEN.sub(_sub, str(text or ""))
  out: List[Tuple[str, str]] = []
  pos = 0
  for m in _SUPERSCRIPT_MARKER.finditer(rendered):
    if m.start() > pos:
      out.append(("text", rendered[pos:m.start()]))
    out.append(("noteref", m.group(1)))
    pos = m.end()
  if pos < len(rendered):
    out.append(("text", rendered[pos:]))
  return out


def build_section_draft_docx(*, business_name: str, run_id: str,
                             section_key: str, payload: Dict[str, Any],
                             cat: FactCatalog,
                             out_dir: Optional[str] = None,
                             now: Optional[_dt.datetime] = None) -> str:
  spec = R.section(section_key)
  now = now or _dt.datetime.now()
  month_year = now.strftime("%B %Y")
  stamp = now.strftime(R.PLAN_FILENAME_STAMP_FORMAT)
  safe_name = _UNSAFE.sub(" ", str(business_name)).strip()
  fname = "%s -- %s -- %s.docx" % (safe_name, spec["title"], stamp)
  out_path = os.path.join(out_dir or R.PLAN_OUTPUT_DIR, fname)

  doc = Document()
  REN._styles(doc)
  _note_ref_style(doc)

  # ---- title page (no header/footer; SECTION DRAFT named out loud)
  doc.add_paragraph(business_name, style="Title")
  doc.add_paragraph("%s — Section Draft" % spec["title"], style="Plan Subtitle")
  doc.add_paragraph("Prepared %s" % month_year, style="Plan Subtitle")

  # ---- the section, under the shell's running header and stamped footer
  s2 = doc.add_section(WD_SECTION.NEW_PAGE)
  REN._header(s2, business_name)
  REN._footer(s2, month_year)
  doc.add_paragraph(spec["title"], style="Heading 1")

  paras: Dict[int, List[Dict[str, Any]]] = {}
  for s in payload.get("sentences") or []:
    paras.setdefault(int(s.get("paragraph") or 1), []).append(s)
  for pno in sorted(paras):
    p = doc.add_paragraph()
    for i, s in enumerate(paras[pno]):
      if i:
        p.add_run(" ")
      for kind, value in _runs(str(s.get("text") or ""), cat):
        r = p.add_run(value)
        if kind == "noteref":
          r.style = doc.styles["Plan Note Ref"]

  notes = payload.get("notes") or []
  if notes:
    doc.add_paragraph(R.NOTES_SECTION_TITLE, style="Heading 1")
    for n in notes:
      p = doc.add_paragraph()
      r = p.add_run(str(n.get("id") or ""))
      r.style = doc.styles["Plan Note Ref"]
      body = FACT_TOKEN.sub(
        lambda m: (cat.get_quiet(m.group(1)).render()
                   if cat.get_quiet(m.group(1)) is not None else m.group(0)),
        str(n.get("text") or ""))
      p.add_run(" %s — %s" % (str(n.get("kind") or ""), body))

  # ---- minimal appendix: the run identifier's ONLY legal home (rule 21)
  doc.add_paragraph("Appendix", style="Heading 1")
  doc.add_paragraph("Run identifier: %s" % run_id, style="Plan Chrome")

  os.makedirs(os.path.dirname(out_path), exist_ok=True)
  doc.save(out_path)
  return out_path
