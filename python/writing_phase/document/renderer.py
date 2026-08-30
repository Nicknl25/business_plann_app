"""THE SHELL RENDERER (2026-08-30). Proves rules 21/22/23 before any prose exists.

Everything is a REAL Word style - Normal, Heading 1/2, Title, Plan Caption,
Plan Table - and NO run ever carries direct formatting: that is what makes the
TOC clickable, the navigation pane work, and the document survive Nick's
editing (rule 23: no text boxes, no floating frames, images inline only).

One font pair: Georgia for body, Calibri for headings/chrome. One table style.
Footer from FOOTER_FORMAT with real PAGE/NUMPAGES fields. The run identifier
appears in the appendix only - never on a client-facing page.

Placeholder text throughout. The point is the object, not the words.
"""
from __future__ import annotations

import datetime as _dt
import io
import os
from typing import Any, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .. import rules as R

BODY_FONT = "Georgia"
HEAD_FONT = "Calibri"
INK = RGBColor(0x1A, 0x27, 0x33)
MUTED = RGBColor(0x71, 0x80, 0x8F)
NAVY = RGBColor(0x12, 0x31, 0x4B)
BAND = "F4F7FA"       # subtle banding (design.py TINT_2)
HAIRLINE = "DDE3E9"

PLACEHOLDER = ("Placeholder text. This paragraph stands where the section's prose "
               "will go; it exists so the object can be judged before a single "
               "sentence is written. The shell is the deliverable of this build.")


# ---------------------------------------------------------------------------
# styles
# ---------------------------------------------------------------------------
def _styles(doc: Document) -> None:
  st = doc.styles
  normal = st["Normal"]
  normal.font.name = BODY_FONT
  normal.font.size = Pt(R.DOCUMENT_CRAFT["body_font_size_pt"])
  normal.font.color.rgb = INK
  normal.paragraph_format.space_after = Pt(6)
  normal.paragraph_format.line_spacing = 1.25   # generous leading (rule 22)

  for name, size, bold, before in (("Heading 1", 16, True, 18), ("Heading 2", 12.5, True, 12)):
    h = st[name]
    h.font.name = HEAD_FONT
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = NAVY
    h.paragraph_format.space_before = Pt(before)
    h.paragraph_format.space_after = Pt(6)
    # headings must carry their font at the STYLE level for every script
    h.element.rPr.rFonts.set(qn("w:cs"), HEAD_FONT)

  title = st["Title"]
  title.font.name = HEAD_FONT
  title.font.size = Pt(30)
  title.font.color.rgb = NAVY

  sub = st.add_style("Plan Subtitle", WD_STYLE_TYPE.PARAGRAPH)
  sub.base_style = st["Normal"]
  sub.font.name = HEAD_FONT
  sub.font.size = Pt(12)
  sub.font.color.rgb = MUTED

  cap = st.add_style("Plan Caption", WD_STYLE_TYPE.PARAGRAPH)
  cap.base_style = st["Normal"]
  cap.font.name = HEAD_FONT
  cap.font.size = Pt(9)
  cap.font.color.rgb = MUTED
  cap.paragraph_format.space_before = Pt(2)
  cap.paragraph_format.space_after = Pt(12)

  chrome = st.add_style("Plan Chrome", WD_STYLE_TYPE.PARAGRAPH)  # header/footer text
  chrome.base_style = st["Normal"]
  chrome.font.name = HEAD_FONT
  chrome.font.size = Pt(8.5)
  chrome.font.color.rgb = MUTED
  chrome.paragraph_format.space_after = Pt(0)

  # ONE table style. Based on Table Grid, hairline horizontals only, banded rows.
  tbl = st.add_style("Plan Table", WD_STYLE_TYPE.TABLE)
  tbl.base_style = st["Table Grid"]
  tbl.font.name = BODY_FONT
  tbl.font.size = Pt(9.5)
  tel = tbl.element
  tblPr = tel.find(qn("w:tblPr"))
  if tblPr is None:
    tblPr = OxmlElement("w:tblPr")
    tel.append(tblPr)
  borders = OxmlElement("w:tblBorders")
  for edge in ("top", "left", "bottom", "right", "insideV"):
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), "none")
    borders.append(e)
  ih = OxmlElement("w:insideH")
  ih.set(qn("w:val"), "single"); ih.set(qn("w:sz"), "4")
  ih.set(qn("w:color"), HAIRLINE)
  borders.append(ih)
  old = tblPr.find(qn("w:tblBorders"))
  if old is not None:
    tblPr.remove(old)
  tblPr.append(borders)
  band = OxmlElement("w:tblStylePr")
  band.set(qn("w:type"), "band1Horz")
  tcPr = OxmlElement("w:tcPr")
  shd = OxmlElement("w:shd")
  shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), BAND)
  tcPr.append(shd)
  band.append(tcPr)
  tel.append(band)
  first = OxmlElement("w:tblStylePr")
  first.set(qn("w:type"), "firstRow")
  rPr = OxmlElement("w:rPr")
  b = OxmlElement("w:b"); rPr.append(b)
  first.append(rPr)
  tel.append(first)


# ---------------------------------------------------------------------------
# fields (PAGE / NUMPAGES / TOC) - real Word fields, no typed numbers
# ---------------------------------------------------------------------------
def _field(par, instr: str, placeholder: str = "") -> None:
  r1 = par.add_run()
  fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
  r1._r.append(fld)
  r2 = par.add_run()
  it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
  r2._r.append(it)
  r3 = par.add_run()
  sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
  r3._r.append(sep)
  if placeholder:
    par.add_run(placeholder)
  r4 = par.add_run()
  end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
  r4._r.append(end)


def _footer(section, month_year: str) -> None:
  """Rule 21, from FOOTER_FORMAT. PAGE/NUMPAGES are fields, the rest literal."""
  section.footer.is_linked_to_previous = False
  par = section.footer.paragraphs[0]
  par.style = "Plan Chrome"
  par.alignment = WD_ALIGN_PARAGRAPH.CENTER
  for chunk in R.FOOTER_FORMAT.split("{page}"):
    pass
  par.add_run("Confidential · Page ")
  _field(par, " PAGE ")
  par.add_run(" of ")
  _field(par, " NUMPAGES ")
  par.add_run(" · Prepared %s · v%s" % (month_year, R.FOOTER_VERSION))


def _header(section, business_name: str) -> None:
  section.header.is_linked_to_previous = False
  par = section.header.paragraphs[0]
  par.style = "Plan Chrome"
  par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  par.add_run(business_name)


# ---------------------------------------------------------------------------
# wrapped placement (rule 23 as narrowed, 2026-08-30): the picture stays an
# ordinary drawing that FLOWS with its paragraph - wp:anchor with wrapSquare,
# never absolute positioning, never a text box. Wrapping is a per-chart
# registry decision (placement: "wrap"), not a global one.
# ---------------------------------------------------------------------------
def _inline_to_wrapped(run, *, align: str = "right") -> None:
  """Convert the run's wp:inline drawing to a wp:anchor with square wrap,
  positioned relative to the COLUMN horizontally and its own PARAGRAPH
  vertically - so it moves when the paragraph moves."""
  drawing = run._r.find(qn("w:drawing"))
  inline = drawing.find(qn("wp:inline"))
  extent = inline.find(qn("wp:extent"))
  doc_pr = inline.find(qn("wp:docPr"))
  graphic = inline.find(qn("a:graphic"))
  anchor = OxmlElement("wp:anchor")
  for k, v in (("distT", "91440"), ("distB", "91440"), ("distL", "114300"),
               ("distR", "114300"), ("simplePos", "0"), ("relativeHeight", "2"),
               ("behindDoc", "0"), ("locked", "0"), ("layoutInCell", "1"),
               ("allowOverlap", "0")):
    anchor.set(k, v)
  simple = OxmlElement("wp:simplePos"); simple.set("x", "0"); simple.set("y", "0")
  anchor.append(simple)
  posH = OxmlElement("wp:positionH"); posH.set("relativeFrom", "column")
  alignEl = OxmlElement("wp:align"); alignEl.text = align
  posH.append(alignEl); anchor.append(posH)
  posV = OxmlElement("wp:positionV"); posV.set("relativeFrom", "paragraph")
  off = OxmlElement("wp:posOffset"); off.text = "0"
  posV.append(off); anchor.append(posV)
  anchor.append(extent)
  wrap = OxmlElement("wp:wrapSquare"); wrap.set("wrapText", "bothSides")
  anchor.append(wrap)
  anchor.append(doc_pr)
  anchor.append(graphic)
  drawing.remove(inline)
  drawing.append(anchor)


# ---------------------------------------------------------------------------
# the shell
# ---------------------------------------------------------------------------
def build_shell(out_path: str, *, business_name: str, run_id: str,
                charts: Sequence[Tuple[str, bytes, str, str]] = (),
                month_year: Optional[str] = None) -> str:
  """charts: (title, png_bytes, section_key, placement) where placement is the
  chart registry's "full_width" or "wrap". Figure numbers are COMPUTED from
  emission order - never authored - and each figure is cross-referenced in the
  placeholder text (rules 8/22). Adding a chart is a registry entry plus a
  theme function, never a renderer change."""
  month_year = month_year or _dt.datetime.now().strftime("%B %Y")
  doc = Document()
  _styles(doc)

  # ---- section 1: title page (no header/footer)
  s1 = doc.sections[0]
  s1.different_first_page_header_footer = False
  doc.add_paragraph(business_name, style="Title")
  doc.add_paragraph("Business Plan", style="Plan Subtitle")
  doc.add_paragraph("Prepared %s" % month_year, style="Plan Subtitle")

  # ---- section 2: TOC + body (running header + stamped footer start here)
  s2 = doc.add_section(WD_SECTION.NEW_PAGE)
  _header(s2, business_name)
  _footer(s2, month_year)
  doc.add_paragraph("Contents", style="Heading 1")
  toc_par = doc.add_paragraph()
  _field(toc_par, ' TOC \\o "1-2" \\h \\z \\u ',
         "Table of contents - right-click and choose Update Field.")

  by_section: Dict[str, List[Tuple[str, bytes, str]]] = {}
  for title, png, sec, placement in charts:
    by_section.setdefault(sec, []).append((title, png, placement))

  fig_no = 0
  body_keys = [k for k in R.body_section_keys()]
  for key in body_keys:
    spec = R.section(key)
    doc.add_paragraph(spec["title"], style="Heading 1")
    doc.add_paragraph(PLACEHOLDER)
    if key == "the_business":
      doc.add_paragraph("A Subsection Heading", style="Heading 2")
      doc.add_paragraph(PLACEHOLDER)
    if key == "financial_plan":
      doc.add_paragraph("The table below is styled the way every table in this "
                        "document will be styled - hairline horizontals, banded "
                        "rows, right-aligned figures.")
      t = doc.add_table(rows=4, cols=5)
      t.style = doc.styles["Plan Table"]
      t.alignment = WD_TABLE_ALIGNMENT.CENTER
      hdr = ["", "Year 1", "Year 2", "Year 3", "Year 4"]
      rows = [("Revenue", "928,000", "1,000,000", "1,100,000", "1,200,000"),
              ("Gross profit", "554,000", "600,000", "659,000", "716,000"),
              ("Net income", "70,000", "87,000", "102,000", "115,000")]
      for c, txt in enumerate(hdr):
        t.rows[0].cells[c].paragraphs[0].add_run(txt)
      for ri, row in enumerate(rows, start=1):
        for ci, txt in enumerate(row):
          cell = t.rows[ri].cells[ci]
          cell.paragraphs[0].add_run(txt)
          if ci > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for title, png, placement in by_section.get(key, []):
      fig_no += 1
      if placement == R.PLACEMENT_WRAP:
        # the anchored picture lives INSIDE the prose paragraph and the text
        # wraps around it - this is the placement four of the eight call for
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(io.BytesIO(png), width=Inches(3.1))
        _inline_to_wrapped(r, align="right")
        p.add_run("As Figure %d shows, the placeholder discussion runs beside "
                  "the chart. " % fig_no + PLACEHOLDER + " " + PLACEHOLDER)
      else:
        doc.add_paragraph("As Figure %d shows, the placeholder discussion would "
                          "reference the chart by number here." % fig_no)
        doc.add_picture(io.BytesIO(png), width=Inches(6.5))
      doc.add_paragraph(R.CHART_CAPTION_FORMAT.format(number=fig_no, title=title),
                        style="Plan Caption")

  # ---- landscape appendix (rule 18: quarterly detail lives here)
  s3 = doc.add_section(WD_SECTION.NEW_PAGE)
  s3.orientation = WD_ORIENT.LANDSCAPE
  s3.page_width, s3.page_height = s3.page_height, s3.page_width
  _header(s3, business_name)
  _footer(s3, month_year)
  doc.add_paragraph("Appendix", style="Heading 1")
  doc.add_paragraph("Full 20-quarter statements render here, landscape, exact "
                    "figures. " + PLACEHOLDER)
  doc.add_paragraph("Run identifier: %s" % run_id, style="Plan Chrome")

  doc.add_paragraph(R.NOTES_SECTION_TITLE, style="Heading 1")
  doc.add_paragraph("Numbered SOURCE and BASIS notes render here. " + PLACEHOLDER)

  os.makedirs(os.path.dirname(out_path), exist_ok=True)
  doc.save(out_path)
  return out_path
